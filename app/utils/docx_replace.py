import re
from docx import Document

def docx_replace(
    doc_obj,
    version: int,
) -> None:
    parapgraphs = [i for i in doc_obj.paragraphs]

    # ИДЕМ ОТ 1 ДО -1 Т.К. ПЕРВЫЙ И ПОСЛЕДНИЙ ВСЕГДА РУТ ЕГО НАДО ПИСАТЬ СТРОГО ТАК!!!
    for i in range(1, len(parapgraphs)-1):
        par_text = parapgraphs[i].text
        access_levels ={
            1: [255, 1, 7, 3, 5],
            2: [255, 2, 7, 3, 6],
            4: [255, 4, 7, 5, 6]
        }
        if ('{{block' in par_text):
            tc = i + 1
            stack = []
            start_index = 0
            start_parapgraph = i
            end_paragraph = 99999
            end_index = 99999
            # ПОИСК ИНДЕкСА ОТКРЫВАЮИЕЙ СКОБКИ
            for j in range(len(par_text)):
                if par_text[j] == '{':
                    start_index = j
                    break

            # ИДКЕМ ПО ТЕКУЩЕМУ ПАРАГРАФУ В НАДЕЖДЕ НАЙТИ БАЛАНС
            for j in range(start_index, len(par_text)):
                if par_text[j] == '{':
                    stack.append('{')
                    continue
                if par_text[j] == '}':
                    stack.pop(0)
                    if len(stack) == 0:
                        end_index = j
                        end_paragraph = i
                        break

            # СРАЗУ СНОСИМ ЕСЛИ В ОДНОМ ПАРАГРАФЕ, ЧТОБЫ НЕ ПАРИТЬСЯ
            if end_paragraph != 99999:

                levels = re.findall(r'(block level)="(\d{1,3})"', par_text)#type: ignore
                levels = [int(i[1]) for i in levels]
                outers_levels = ['{{block level' + f'="{i}"' for i in levels]
                outers = []

                for level in range(len(outers_levels) - 1):
                    outers.append(par_text[par_text.find(outers_levels[level]): par_text.find(outers_levels[level + 1])])
                outers.append(par_text[par_text.find(outers_levels[-1]): len(par_text)])

                indexs = []
                i = 0
                text = parapgraphs[start_parapgraph].text
                for level in levels:
                    if level not in access_levels[version]:
                        start_index = text.find(outers[i])
                        if i != len(outers) - 1:
                            end_index = text.find(outers[i + 1])
                        else:
                            end_index = len(text)
                        indexs.append([start_index, end_index])
                    i += 1
                if indexs == []: pass    
                else:
                    if len(indexs) == 1:
                        parapgraphs[start_parapgraph].text = \
                            parapgraphs[start_parapgraph].text[:indexs[0][0]] \
                                + parapgraphs[start_parapgraph].text[indexs[0][0]+indexs[0][1]:]
                    else:
                        le = 0
                        for el in indexs:
                            s = el[0] - le
                            e = el[1] - le
                            le = e - s 
                            parapgraphs[start_parapgraph].text = \
                                parapgraphs[start_parapgraph].text[:s] +\
                                    parapgraphs[start_parapgraph].text[e:]  
                            
                        data = re.findall(r'[|](\w{5,})}}}}', parapgraphs[start_parapgraph].text)
                        data_2 = re.findall(r'{{block\s*level=\"\d{1,3}\"\|\n*[^}]+}}.?}}', parapgraphs[start_parapgraph].text)
                        index_data = parapgraphs[start_parapgraph].text.find(data_2[0])
                        parapgraphs[start_parapgraph].text = parapgraphs[start_parapgraph].text[:index_data] + \
                            data[0] + parapgraphs[start_parapgraph].text[index_data+len(data_2[0]):]
                continue

            # БАЛАНС НЕ НАЙДЕН, НАЧИНАЕМ ХОДИТЬ ПО ДРУГИМ ПАРАГРАФАМ
            if end_paragraph == 99999:
                while len(stack) != 0:
                    next_par_text = parapgraphs[tc].text
                    for j in range(0, len(next_par_text)):
                        if next_par_text[j] == '{':
                            stack.append('{')
                            continue
                        if next_par_text[j] == '}':
                            stack.pop(0)
                            if len(stack) == 0:
                                end_index = j
                                end_paragraph = tc
                                break
                    tc += 1

            # СНОСИМ!!! 🤗🤗🤗😊😊😊😊😊😊😁😁😁👍👍
            levels = re.findall(r'(block level)="(\d)"', par_text)#type: ignore
            levels = [int(i[1]) for i in levels]
            flag = False
            for level in levels:
                if level not in levels:
                    flag = True
                    break
            
            if flag:
                for k in range(start_parapgraph, end_paragraph+1):
                    if k == start_parapgraph:
                        parapgraphs[k].text = parapgraphs[k].text[:start_index]
                    elif k == end_paragraph:
                        parapgraphs[k].text = parapgraphs[k].text[end_index+1:]
                    else:
                        parapgraphs[k].text = ""

document_1 = Document("./vvit.docx")
docx_replace(document_1, 1)
document_1.save(f"./vvit_1.docx")

document_2 = Document("./vvit.docx")
docx_replace(document_2, 2)
document_2.save(f"./vvit_2.docx")

document_4 = Document("./vvit.docx")
docx_replace(document_4, 4)
document_4.save(f"./vvit_4.docx")