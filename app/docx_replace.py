import re
from docx import Document


def docx_replace(
    doc_obj,
    version: int,
) -> None:
    parapgraphs = [i for i in doc_obj.paragraphs]

    for i in range(0, len(parapgraphs)):
        par_text = parapgraphs[i]
        runs = par_text.runs
        for r in runs:
            print(r.text)


document_1 = Document("./vvit.docx")
docx_replace(document_1, 1)
# document_1.save(f"./vvit_1.docx")

# document_2 = Document("./vvit.docx")
# docx_replace(document_2, 2)
# document_2.save(f"./vvit_2.docx")

# document_4 = Document("./vvit.docx")
# docx_replace(document_4, 4)
# document_4.save(f"./vvit_4.docx")
